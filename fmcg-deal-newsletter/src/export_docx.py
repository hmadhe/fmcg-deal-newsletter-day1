"""
Renders the structured newsletter dict (from newsletter_gen.build_newsletter)
into a formatted Word document using python-docx.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BRAND_COLOR = RGBColor(0x1F, 0x4E, 0x79)
BADGE_COLOR = RGBColor(0xC9, 0xDA, 0xEE)


def _set_cell_background(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _remove_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _set_cell_margins(cell, top=200, start=300, bottom=200, end=300):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def _add_hero_banner(doc, newsletter):
    """Title block styled as a solid-color banner: date badge, title, meta line."""
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    _set_cell_background(cell, "1F4E79")
    _set_cell_margins(cell)

    badge_p = cell.paragraphs[0]
    badge_run = badge_p.add_run(newsletter["period"].upper())
    badge_run.font.size = Pt(9)
    badge_run.font.bold = True
    badge_run.font.color.rgb = BADGE_COLOR

    title_p = cell.add_paragraph()
    title_run = title_p.add_run(newsletter["title"])
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    meta_p = cell.add_paragraph()
    meta_run = meta_p.add_run(
        f"Generated {newsletter['generated_on']}  |  {newsletter['total_deals_found']} deals tracked this period"
    )
    meta_run.italic = True
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = BADGE_COLOR

    doc.add_paragraph()


def export_newsletter_to_docx(newsletter, output_path="../output/newsletter.docx"):
    doc = Document()

    _add_hero_banner(doc, newsletter)

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(newsletter["executive_summary"])

    doc.add_heading(f"Top Deals This Period ({len(newsletter['top_deals'])})", level=1)
    for deal in newsletter["top_deals"]:
        _add_deal_block(doc, deal, heading_level=2)

    if newsletter["other_deals"]:
        doc.add_heading("Other Notable Activity", level=1)
        for deal in newsletter["other_deals"]:
            _add_deal_block(doc, deal, heading_level=3)

    doc.add_page_break()
    doc.add_heading("Methodology", level=1)
    doc.add_paragraph(
        f"Generated from {newsletter['total_deals_found']} deduplicated, relevance- and "
        "credibility-scored FMCG news articles, refined using Gemini-based structured "
        "extraction. Deal facts are taken directly from extraction output; only the "
        "executive summary above is LLM-generated free-form prose. Each deal card links "
        "back to its original source article for verification."
    )

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def _add_deal_block(doc, deal, heading_level=2):
    h = doc.add_heading(deal["headline"], level=heading_level)
    for run in h.runs:
        run.font.color.rgb = BRAND_COLOR

    p = doc.add_paragraph()
    p.add_run(f"{deal['acquirer']} \u2192 {deal['target']}").bold = True
    p.add_run(f"   |   {deal['deal_type'].replace('_', ' ').title()}   |   {deal['deal_value']}")

    doc.add_paragraph(deal["summary"])

    n = deal["corroboration_count"]
    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run(
        f"Deal ID: {deal['deal_id']}   |   Sources: {', '.join(deal['sources'])}   |   "
        f"Corroboration: {n} independent source{'s' if n != 1 else ''}   |   "
        f"Confidence: {deal['confidence']}"
    )
    meta_run.italic = True
    meta_run.font.size = Pt(9)


if __name__ == "__main__":
    import json
    with open("../data/newsletter.json", "r", encoding="utf-8") as f:
        newsletter = json.load(f)
    path = export_newsletter_to_docx(newsletter)
    print(f"Saved -> {path}")
